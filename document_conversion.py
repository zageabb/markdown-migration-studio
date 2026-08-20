from __future__ import annotations

from pathlib import Path
import re


def _markdown_escape(value: object) -> str:
    return str(value or "").replace("\n", " ").strip().replace("|", "\\|")


def _rows_to_markdown(rows: list[list[str]]) -> str:
    cleaned = [[_markdown_escape(cell) for cell in row] for row in rows if any(str(cell).strip() for cell in row)]
    if not cleaned:
        return ""
    width = max(len(row) for row in cleaned)
    padded = [row + [""] * (width - len(row)) for row in cleaned]
    return "\n".join([
        "| " + " | ".join(padded[0]) + " |",
        "| " + " | ".join(["---"] * width) + " |",
        *["| " + " | ".join(row) + " |" for row in padded[1:]],
    ])


def docx_to_markdown(path: str | Path) -> str:
    """Convert a DOCX file using the shared converter behavior from the Ollama apps."""
    from docx import Document

    document = Document(str(path))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name or "").lower()
        if style.startswith("heading"):
            digits = "".join(character for character in style if character.isdigit())
            level = min(max(int(digits or "1"), 1), 6)
            lines.append(f"{'#' * level} {text}")
        elif "list bullet" in style:
            lines.append(f"- {text}")
        elif "list number" in style:
            lines.append(f"1. {text}")
        else:
            lines.append(text)

    for table in document.tables:
        rendered = _rows_to_markdown([[cell.text.strip() for cell in row.cells] for row in table.rows])
        if rendered:
            lines.extend(["", rendered])

    markdown = "\n".join(lines).strip()
    if not markdown:
        raise ValueError("No readable text could be extracted from the Word document")
    return markdown + "\n"


def markdown_to_docx(markdown: str, destination: str | Path) -> None:
    """Write a readable Word copy of the Markdown produced by the app."""
    from docx import Document
    from docx.shared import Pt

    document = Document()
    lines = markdown.splitlines()
    index = 0
    in_code_block = False
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            index += 1
            continue
        if in_code_block:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            index += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            document.add_heading(heading.group(2), level=len(heading.group(1)))
            index += 1
            continue
        if "|" in line and index + 1 < len(lines) and _is_table_separator(lines[index + 1]):
            rows = [_markdown_table_cells(line)]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                rows.append(_markdown_table_cells(lines[index]))
                index += 1
            width = max(len(row) for row in rows)
            table = document.add_table(rows=len(rows), cols=width)
            table.style = "Table Grid"
            for row_number, row in enumerate(rows):
                for column_number, value in enumerate(row):
                    table.cell(row_number, column_number).text = value
            continue
        bullet = re.match(r"^[-*+]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if bullet:
            document.add_paragraph(bullet.group(1), style="List Bullet")
        elif numbered:
            document.add_paragraph(numbered.group(1), style="List Number")
        elif stripped:
            document.add_paragraph(stripped)
        index += 1

    destination = Path(destination)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))


def _markdown_table_cells(line: str) -> list[str]:
    return [cell.strip().replace("\\|", "|") for cell in line.strip().strip("|").split("|")]


def _is_table_separator(line: str) -> bool:
    cells = _markdown_table_cells(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)
